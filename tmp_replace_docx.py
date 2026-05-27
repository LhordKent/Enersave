from docx import Document
from pathlib import Path

in_path = Path(r"C:\Users\Admin\Downloads\FR1-Template.docx")
out_path = Path(r"C:\Users\Admin\Downloads\FR1-Template-Enersave.docx")
doc = Document(str(in_path))

repl = [
    ("DevoutNet", "Enersave"),
    ("CHURCH MANAGEMENT SYSTEM", "ENERGY OPERATIONS PLATFORM"),
    ("church management system", "building energy operations platform"),
    ("church admin", "facility manager"),
    ("Church Admin", "Facility Manager"),
    ("churchgoer", "operations user"),
    ("Churchgoers", "Operations Users"),
    ("churchgoers", "operations users"),
    ("Churchgoer", "Operations User"),
    ("system admin", "platform admin"),
    ("System Admin", "Platform Admin"),
    ("churches", "buildings"),
    ("church", "building"),
    ("Church", "Building"),
    ("donations", "energy records"),
    ("Donations", "Energy Records"),
    ("prayer requests", "service requests"),
    ("Prayer requests", "Service requests"),
    ("prayer", "service"),
    ("Prayer", "Service"),
    ("certificate", "report"),
    ("Certificate", "Report"),
    ("requests", "control requests"),
    ("Requests", "Control Requests"),
    ("Payment Gateway", "Integration Gateway"),
    ("payment gateway", "integration gateway"),
    ("CALENDAR OF EVENTS", "ENERGY EVENTS"),
    ("Calendar of Events", "Energy Events"),
    ("User Interface Design displays the visual representation of the proposed Enersave.",
     "User Interface Design displays the visual representation of the proposed Enersave Energy Operations platform."),
    ("This web-based building energy operations platform focuses on the functionality and user experience to satisfy them with ease of usability.",
     "This web-based platform focuses on live telemetry, room controls, AI insights, alerts, and reports with a user-friendly experience."),
    ("Software Specification", "Software Specification (Enersave)"),
    ("Hardware Specification", "Hardware Specification (Enersave)"),
    ("Html 5, CSS 3, JavaScript, bootstrap, and PHP were used to develop a friendly user interface.  To retrieve information, MySQL is used as a database.",
     "HTML, CSS, JavaScript, and a modern React/Next.js interface are used to deliver a responsive dashboard experience, with backend data services for telemetry and reporting."),
]

# High-level exact replacements
for p in doc.paragraphs:
    if p.text.strip() == "DevoutNet: A WEB-BASED":
        p.text = "Enersave: A WEB-BASED"
    elif p.text.strip() == "CHURCH MANAGEMENT SYSTEM":
        p.text = "ENERGY OPERATIONS PLATFORM"

count = 0

def replace_text(t: str):
    global count
    nt = t
    for a,b in repl:
        if a in nt:
            c = nt.count(a)
            nt = nt.replace(a,b)
            count += c
    return nt

for p in doc.paragraphs:
    new_text = replace_text(p.text)
    if new_text != p.text:
        p.text = new_text

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            new_text = replace_text(cell.text)
            if new_text != cell.text:
                cell.text = new_text

# Tune first summary/title paragraphs if present
for p in doc.paragraphs:
    if p.text.strip().startswith("Figure 1 shows the home page of Enersave"):
        p.text = ("Figure 1 shows the Enersave home page where users can monitor live building energy data, room controls, alerts, reports, and AI-driven insights.")

# Add concise Enersave platform statement near software spec if line exists
for i,p in enumerate(doc.paragraphs):
    if p.text.strip() == "Software Specification (Enersave)":
        if i + 1 < len(doc.paragraphs):
            doc.paragraphs[i+1].text = "This section summarizes the software environment used by Enersave Energy Operations."

# Save
doc.save(str(out_path))
print(f"Saved: {out_path}")
print(f"Total replacements: {count}")
