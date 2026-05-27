from docx import Document
p = r"C:\Users\Admin\Downloads\FR1-Template.docx"
d = Document(p)
for i, para in enumerate(d.paragraphs, 1):
    t = para.text.strip()
    if t:
        print(f"P{i}: {t}")
print("--TABLES--")
for ti, table in enumerate(d.tables, 1):
    print(f"Table {ti}")
    for ri, row in enumerate(table.rows, 1):
        vals = [c.text.strip().replace("\n", " | ") for c in row.cells]
        print(f" R{ri}: " + " || ".join(vals))
